import os
import re
import copy


from docx import Document
from stream import ClassStream
from error import ClassError


class ClassSentencePath:
    """句子地址对象
        通过句子输入绝对位置，输入相对位置，以及上个句子的输出绝对位置确定当前句子的输出位置
    """

    LEN_TABLE = 4
    LEN_PARA = 2

    LAST_PATH_TABLE_WRITE = [0, 0, 0, 0]
    LAST_PATH_TABLE_READ = [0, 0, 0, 0]
    INTERVAL_TABLE = False
    INTERVAL_PARA = False
    LAST_PATH_PARA_WRITE = [0, 0]
    LAST_PATH_PARA_READ = [0, 0]

    def __init__(self, path: list):
        """录入绝对位置，计算、储存相对位置"""
        self.Path = copy.deepcopy(path)
        lens = len(self.Path)
        if lens == ClassSentencePath.LEN_TABLE:
            self.LastPath = copy.deepcopy(ClassSentencePath.LAST_PATH_TABLE_READ)
            ClassSentencePath.LAST_PATH_TABLE_READ = copy.deepcopy(path)
        elif lens == ClassSentencePath.LEN_PARA:
            self.LastPath = copy.deepcopy(ClassSentencePath.LAST_PATH_PARA_READ)
            ClassSentencePath.LAST_PATH_PARA_READ = copy.deepcopy(path)

    def path(self):
        """DataFile.fill_in_data调用输出时求地址"""
        lens = len(self.Path)
        if lens == ClassSentencePath.LEN_TABLE:
            return self.__table_path()
        elif lens == ClassSentencePath.LEN_PARA:
            return self.__para_path()

    def __table_path(self):
        if self.LastPath == ClassSentencePath.LAST_PATH_TABLE_WRITE:
            ClassSentencePath.INTERVAL_TABLE = False
            ClassSentencePath.LAST_PATH_TABLE_WRITE = copy.deepcopy(self.Path)
        else:
            ClassSentencePath.INTERVAL_TABLE = True
            interval = [y - x for x, y in zip(self.LastPath, self.Path)]
            if interval[1] < 0:
                interval[1] = 0
            if interval[2] < 0:
                interval[2] = 0
            if interval[3] < 0:
                interval[3] = 0

            if interval[0] != 0:
                ClassSentencePath.LAST_PATH_TABLE_WRITE = copy.deepcopy(self.Path)
            elif interval.count(0) == 3:
                jian = ClassSentencePath.LAST_PATH_TABLE_WRITE
                ClassSentencePath.LAST_PATH_TABLE_WRITE = [x+y for x, y in zip(jian, interval)]
            else:
                ClassError.warn("警告，与上个填入不在同一列或行")
                ClassSentencePath.LAST_PATH_TABLE_WRITE = copy.deepcopy(self.Path)

        return ClassSentencePath.LAST_PATH_TABLE_WRITE

    def __para_path(self):
        if self.LastPath == ClassSentencePath.LAST_PATH_PARA_WRITE:
            ClassSentencePath.INTERVAL_PARA = False
            ClassSentencePath.LAST_PATH_PARA_WRITE = copy.deepcopy(self.Path)
        else:
            ClassSentencePath.INTERVAL_PARA = True
            interval = [y - x for x, y in zip(self.LastPath, self.Path)]
            if interval[0] != 0:
                ClassSentencePath.LAST_PATH_PARA_WRITE = copy.deepcopy(self.Path)
            elif interval.count(0) == 2:
                jian = ClassSentencePath.LAST_PATH_PARA_WRITE
                ClassSentencePath.LAST_PATH_PARA_WRITE = [x + y for x, y in zip(jian, interval)]
            else:
                ClassError.warn("警告，与上个填入不在同一列或行")
                ClassSentencePath.LAST_PATH_PARA_WRITE = copy.deepcopy(self.Path)

        return ClassSentencePath.LAST_PATH_PARA_WRITE

    @staticmethod
    def fresh(table=[0, 0, 0, 0], para=[0, 0]):
        """刷新对象间信息"""
        ClassSentencePath.LAST_PATH_TABLE_WRITE = table
        ClassSentencePath.LAST_PATH_PARA_WRITE = para
        ClassSentencePath.INTERVAL_TABLE = False
        ClassSentencePath.INTERVAL_PARA = False


class ClassSentences:
    """处理中替换句子

        处理中替换句子的表达方式

        :param source:   该句的位置para，table，path (of output docx)，name。用于决定回溯生成的操作位置
        :type source:    str
        :param path:     该句的来源位置，包含相对与绝对位置，用于决定填入文本时的位置
        :type path:      list
        :param text:     替换句子的原句
        :type text:      str
        :param index:    替换句子中的替换变量列表
    """
    def __init__(self, source: str, path: list, text: str):
        self.Source = source
        self.Path = ClassSentencePath(path)
        self.Text = text
        self.Index = []
        self.Result = None


class ClassDocxFile:
    """处理中文本结构

        处理中对转换的docx的表达

        :param file:        docx文件本体
        :param path:        docx文件输出路径
        :param data_docx:   docx对应替换的数据的存储
        :param data_xlsx:   docx对应xlsx数据的存储
    """
    def __init__(self, file: Document()):
        self.File = file
        self.Path = None
        self.Name = None
        self.DataDocx = None
        self.DataXlsx = None
        self.PositionOffset = []

    def up_file(self, file: Document()):
        """DocxTemplateWithData.__init__调用"""
        self.File = file

    def get_template(self, temp, cells):
        """MainProcessor.data_filling调用复制后初始化

        :param temp: processor.ClassDocxData
        :param cells: xlsxmethod.ClassData
        :return:
        """
        self.DataDocx = temp
        self.DataXlsx = cells
        self.DataDocx.name_init(cells)

    def get_data(self, data: list):
        """MainProcessor。data_filling调用，运行时获取信息"""
        self.DataXlsx.get_data(data)
        self.DataDocx.name_init(data)
        self.DataDocx.get_data(self.DataXlsx)

    def finish_data_input(self):
        """MainProcessor。data_filling调用，输入结束更新最终数据"""
        self.DataXlsx.fresh()

        name, file = self.DataDocx.finish_data_input(self.DataXlsx)
        self.Name = name
        self.Path = file

    def fill_in_docx(self):
        """DocxTemplateWithStream.fill_in_docx调用Docx模板填入
            注意！！！cell中应是paragraphs对象，但由于疏忽以及时间有献未来得及修改
        """
        ClassSentencePath.fresh()
        for enum in self.DataDocx.EnumBodyStream.values_():
            for inputBodyStream in enum.values_():
                for sentence in inputBodyStream.values_():
                    path = sentence.Path.path()
                    lens = len(path)
                    if lens == ClassSentencePath.LEN_TABLE:
                        self.__fill_table_cell(sentence, path)
                    elif lens == ClassSentencePath.LEN_PARA:
                        self.__fill_para(sentence, path)

    def __fill_table_cell(self, sentence, path):
        cell = self.File.tables[path[0]].cell(path[1], path[2])

        while True:
            try:
                para = cell.paragraphs[path[3]]
                break
            except IndexError:
                cell.add_paragraph("")

        para = cell.paragraphs[path[3]]
        text = para.text
        if ClassSentencePath.INTERVAL_TABLE:  # 非漂移必定替换
            text = sentence.Result
        else:
            text = text.replace(sentence.Text, sentence.Result, 1)
        para.text = ""
        # cell.paragraphs[0].clear()
        para.add_run(text)

    def __fill_para(self, sentence, path):
        para = self.File.paragraphs[path[0]]
        text = para.text
        if ClassSentencePath.INTERVAL_PARA:
            text = sentence.Result
        else:
            text = text.replace(sentence.Text, sentence.Result, 1)
        para.text = text

    def output(self, path_list):
        """DocxTemplateWithStream.output调用，结果Docx文件输出"""
        try:
            self.File.save(self.Path + self.Name)
        except FileNotFoundError:
            os.mkdir(self.Path)
            path_list += [self.Path]
            self.File.save(self.Path + self.Name)


class ClassDocxStreamWithTemplate:
    """Docx模板本体及其数据列表
        对docx的预处理及缓存
    """

    def __init__(self, sz_name_template: str, re_test: str = r"(#@.*?@#)"):
        # 包含文件本体及其所有生成文件
        self.DocxStream = ClassStream()

        # mo版文件名
        self.DocxTemplateName = sz_name_template

        # 包含所有处理规则
        self.TemplateSentenceList = []

        # docx模板处理结果 由Processor赋予
        self.DocxTemplateData = None

        self.PureTemplate = ClassDocxFile(Document(sz_name_template))

        # 用相对坐标，排除偏移的影响
        now_para = 0
        la_para = 0
        for para in self.PureTemplate.File.paragraphs:
            replaces = re.findall(re_test, para.text)
            # replaces = list(set(replaces))  # 优化合并同类项
            if len(replaces) == 0:
                now_para += 1
                continue
            else:
                for replace_text in replaces:

                    self.TemplateSentenceList += [ClassSentences("para",  # error with
                                                                 [now_para, now_para - la_para],
                                                                 replace_text)]
                la_para = now_para
                now_para += 1

        now_table = 0
        for table in self.PureTemplate.File.tables:
            now_row = 0
            for row in table.rows:
                now_cell = 0
                for cell in row.cells:
                    now_para = 0
                    if cell.text[-3:] == "$%&":  # 针对不规则表格的优化
                        now_cell += 1
                        continue

                    for para in cell.paragraphs:

                        replaces = re.findall(re_test, para.text)
                        replaces = list(set(replaces))  # 优化合并同类项?

                        if len(replaces) == 0:
                            now_para += 1
                            continue
                        else:
                            for replace_text in replaces:
                                self.TemplateSentenceList += [ClassSentences("table",  # error with
                                                                             [now_table, now_row, now_cell, now_para],
                                                                             replace_text)]

                        now_para += 1

                    cell.text += "$%&"
                    now_cell += 1
                now_row += 1
            now_table += 1

        self.PureTemplate.up_file(Document(sz_name_template))  # 刷新所有表格

    def fill_in_docx(self):
        """使所有模板副本填入数据"""
        for key in self.DocxStream.keys():
            self.DocxStream[key].fill_in_docx()

    def output(self, path_list):
        """结果输出"""
        for key in self.DocxStream.keys():
            self.DocxStream[key].output(path_list)


if __name__ == "__main__":  # test
    docx = ClassDocxStreamWithTemplate("./data/table2.docx")

    for jim in docx.TemplateSentenceList:
        print(jim.Text)
